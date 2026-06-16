import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_text_black(run, font_name="Times New Roman", font_size=12, bold=False):
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.highlight_color = None
    run.font.underline = False

doc_path = r"D:\Kuliah Praktik\KP BRIN\docs\Review_GapAnalysis_AirQuality_Filled.docx"
try:
    doc = docx.Document(doc_path)
except Exception as e:
    print("Gagal membuka dokumen:", e)
    exit(1)

# Tambahkan Judul Bagian Baru
doc.add_page_break()
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("TINJAUAN LITERATUR TEMATIK DAN ANALISIS KESENJANGAN (GAP ANALYSIS)")
set_text_black(run_title, font_size=14, bold=True)

# Paragraf Pembuka
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run1 = p1.add_run(
    "Ulasan komprehensif terhadap 22 artikel penelitian terkini mengenai prediksi kualitas udara "
    "mengungkapkan pergeseran metodologis yang signifikan serta menyoroti beberapa celah penelitian yang masih krusial. "
    "Tinjauan literatur ini diklasifikasikan ke dalam tiga tema utama untuk mensintesis kontribusi dan keterbatasan "
    "studi terdahulu, yang kemudian menjadi landasan rumusan gap analysis dalam penelitian ini."
)
set_text_black(run1)

# Tema A
p2_title = doc.add_paragraph()
run2_title = p2_title.add_run("A. Evolusi dan Performa Algoritma (Deep Learning vs Machine Learning Klasik)")
set_text_black(run2_title, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run2 = p2.add_run(
    "Mayoritas literatur terkini telah meninggalkan model regresi linear sederhana, beralih pada arsitektur "
    "Machine Learning (seperti Random Forest dan XGBoost) serta Deep Learning (seperti LSTM dan CNN-LSTM). "
    "Penelitian menunjukkan bahwa model ansambel seperti XGBoost sangat optimal untuk data tabular berdimensi rendah, "
    "sedangkan arsitektur hybrid Spatio-Temporal seperti CNN-LSTM mendominasi akurasi prediksi untuk dataset yang "
    "melibatkan sebaran spasial stasiun observasi yang luas. Meski demikian, ditemukan kritik mendasar pada beberapa "
    "studi terdahulu yang tidak menggunakan fungsi aktivasi yang tepat untuk regresi atau memiliki rentang data "
    "pelatihan yang terlalu singkat (kurang dari 1 tahun), sehingga rentan terhadap fenomena overfitting serta "
    "gagal menangkap variabilitas cuaca lintas musim."
)
set_text_black(run2)

# Tema B
p3_title = doc.add_paragraph()
run3_title = p3_title.add_run("B. Dimensi Spatio-Temporal, Integrasi Cuaca, dan Pemanfaatan Low-Cost Sensor")
set_text_black(run3_title, bold=True)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run3 = p3.add_run(
    "Penelitian yang paling berhasil mencapai tingkat presisi tinggi adalah studi yang secara inheren mengintegrasikan "
    "variabel meteorologi (seperti suhu, kelembaban, dan profil angin dari ERA5) sebagai fitur eksogen, bukan hanya "
    "mengandalkan data historis PM2.5. Namun, kelangkaan stasiun pemantauan resmi (ground-level stations) sering "
    "menjadi kendala spasial yang signifikan, khususnya di wilayah kepulauan tropis seperti Indonesia. Oleh karena itu, "
    "beberapa studi terbaru mulai memanfaatkan data satelit (Google Earth Engine) maupun jaringan sensor biaya rendah "
    "(low-cost sensor) yang berbasis Internet of Things (IoT) untuk meningkatkan resolusi spasial data udara, meskipun "
    "sensor-sensor ini masih terhambat oleh masalah overestimasi akibat tingkat kelembaban tropis yang tinggi."
)
set_text_black(run3)

# Tema C
p4_title = doc.add_paragraph()
run4_title = p4_title.add_run("C. Kebutuhan Kritis akan Interpretabilitas (Explainable AI / XAI)")
set_text_black(run4_title, bold=True)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run4 = p4.add_run(
    "Kelemahan paling fundamental dari tingginya akurasi model deep learning modern adalah sifatnya yang menyerupai "
    "kotak hitam (black-box). Kondisi ini menyebabkan luaran model sulit diterjemahkan menjadi intervensi kebijakan "
    "yang konkret oleh otoritas lingkungan. Merespons tantangan ini, beberapa penelitian pionir mulai mengintegrasikan "
    "metode Explainable AI, khususnya SHAP (SHapley Additive exPlanations), untuk mengurai kontribusi individual tiap "
    "variabel meteorologis terhadap konsentrasi polutan. Meski demikian, penerapan XAI pada prediksi kualitas udara di "
    "kota-kota besar Indonesia selain Jakarta masih tergolong langka."
)
set_text_black(run4)

# Kesimpulan Gap Analysis
p5_title = doc.add_paragraph()
p5_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5_title = p5_title.add_run("KESIMPULAN GAP ANALYSIS")
set_text_black(run5_title, font_size=12, bold=True)

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run5 = p5.add_run(
    "Berdasarkan sintesis literatur tersebut, kesenjangan (gap) utama penelitian ini dikonfirmasi mencakup tiga aspek "
    "fundamental: (1) Gap Metodologis, yakni mendesaknya kebutuhan arsitektur prediksi yang menggabungkan akurasi "
    "tinggi dari model kompleks dengan transparansi inferensi via Explainable AI; (2) Gap Geografis, dengan "
    "dominannya riset di dataran Tiongkok dan minimnya pemodelan prediksi multi-variabel untuk kawasan tropis non-ibukota "
    "(seperti Jawa Barat/Bandung); dan (3) Gap Data, di mana integrasi fusi antara data sensor biaya rendah dan "
    "parameter meteorologis makro masih belum dieksplorasi secara paripurna dalam satu kerangka modeling."
)
set_text_black(run5)

doc.save(doc_path)
print("Berhasil memasukkan Tinjauan Pustaka Tematik ke DOCX dengan format yang sesuai.")
