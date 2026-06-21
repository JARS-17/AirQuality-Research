import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_report(output_path: str):
    """Create an academic DOCX report with the required sections."""
    doc = Document()
    
    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title (centered)
    title = doc.add_heading('Laporan Ringkas – Analisis Gap & Progres Pengumpulan Data ERA5', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Latar Belakang
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('Latar Belakang\n').bold = True
    p.add_run('Penelitian ini menitikberatkan pada prediksi konsentrasi PM2.5 di Cekungan Bandung menggunakan data iklim reanalisis ERA5 (European Centre for Medium-Range Weather Forecasts) serta data kualitas udara. Tujuan utamanya adalah mengembangkan model prediktif berbasis machine learning yang andal untuk mendukung pemantauan dan kebijakan penanggulangan polusi udara yang berdampak pada kesehatan masyarakat dan lingkungan.')

    # 2. Literatur Review
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nLiteratur Review\n').bold = True
    p.add_run('- Mayoritas studi prediksi PM2.5 menggunakan machine learning (seperti LSTM, Random Forest, XGBoost) masih berfokus pada wilayah non-tropis atau negara empat musim seperti Tiongkok, India, Eropa, dan Amerika Serikat. Penelitian di kawasan Asia Tenggara, khususnya di Indonesia yang beriklim tropis, masih sangat minim.\n- Model-model konvensional umumnya hanya memanfaatkan variabel dasar seperti suhu, kelembapan, dan kecepatan angin. Variabel-variabel kompleks seperti lapisan batas planet (boundary layer height), kandungan ozon (total column ozone), dan tutupan awan (total cloud cover) masih jarang dieksplorasi mendalam.\n- Implementasi teknik interpretabilitas model (Explainable AI / XAI) seperti SHAP atau LIME pada prediksi PM2.5 di daerah tropis belum banyak diterapkan, sehingga insight bagi perumusan kebijakan masih belum optimal.')

    # 3. Analisis Gap & Pertanyaan Riset
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nAnalisis Gap\n').bold = True
    p.add_run('1. Keterbatasan studi prediksi AQI / PM2.5 yang komprehensif di Indonesia menggunakan data observasi grid (ERA5) yang beresolusi tinggi.\n2. Variabel meteorologis yang kurang dimanfaatkan, padahal variabel kompleks (lapisan batas, ozon) secara teoritis memiliki korelasi kuat terhadap akumulasi partikulat PM2.5.\n3. Kurangnya interpretabilitas, menjadikan model AI sebagai "black-box" yang sulit diinterpretasikan penyebab dominannya.')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nPertanyaan Riset\n').bold = True
    p.add_run('1. Sejauh mana penggunaan data beresolusi tinggi dan variabel komprehensif dari ERA5 dapat meningkatkan akurasi prediksi PM2.5 di wilayah Cekungan Bandung?\n2. Bagaimana kontribusi dan pengaruh signifikan dari variabel-variabel meteorologis (suhu, curah hujan, lapisan batas, hingga ozon) terhadap dinamika PM2.5 di Cekungan Bandung?\n3. Bagaimana penerapan model Explainable AI (XAI) seperti SHAP dapat mengidentifikasi penggerak utama polusi udara secara temporer maupun musiman?')

    # 4. Pemahaman Dataset ERA5
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nPemahaman Dataset ERA5\n').bold = True
    p.add_run('- Dataset: Copernicus CDS – Reanalysis ERA5 Single Levels\n- Periode: Januari 2022 hingga Desember 2026\n- Resolusi Spasial: 0.25° × 0.25° yang mencakup wilayah Cekungan Bandung (Batas Koordinat: N -6.80, W 107.50, S -7.05, E 107.75)\n- Cakupan Waktu: Data observasi per jam (24 jam sehari) secara lengkap\n- Variabel Terpilih (10 Variabel): 10m u-component of wind, 10m v-component of wind, 2m dewpoint temperature, 2m temperature, surface pressure, boundary layer height, total column ozone, total column water vapour, total cloud cover, dan total precipitation.\n- Progres Saat Ini: Sedang berlangsungnya proses otomasi pengunduhan menggunakan skrip Python untuk 30 batch (per 2 bulan) untuk menghindari batasan sistem.')

    # 5. Dataset yang Diperlukan (Tambahan)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nDataset yang Diperlukan (Tambahan)\n').bold = True
    p.add_run('- Data Target Kualitas Udara (PM2.5): Data historis konsentrasi PM2.5 per jam dari OpenAQ, ISPUnet (KLHK), atau jaringan low-cost sensor di sekitar Bandung.\n- Data Topografi/Geografis (Opsional): Data digital elevation model (DEM) untuk menambah fitur ketinggian jika diperlukan.\n- Data Satelit Pembanding (Opsional): Data penginderaan jauh tambahan untuk validasi iklim secara mikro.')

    # 6. Apa yang Belum Dikerjakan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nApa yang Belum Dikerjakan\n').bold = True
    p.add_run('- Penyelesaian unduhan keseluruhan batch data ERA5.\n- Pengumpulan dan pembersihan (preprocessing) data historis PM2.5 sebagai target prediksi.\n- Feature engineering: penggabungan data metereologis (ERA5) dan PM2.5, serta pembuatan fitur tambahan seperti time-lags dan agregasi statistik.\n- Pengembangan dan pengujian model (baseline RF/XGBoost, serta deep learning seperti LSTM/BiLSTM).\n- Implementasi SHAP/LIME untuk mendapatkan interpretasi variabel prediktif.\n- Penulisan manuskrip akhir/draft hasil.')

    # 7. Perencanaan Penelitian
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nPerencanaan Penelitian\n').bold = True
    p.add_run('- Tahap 1 (Bulan 1): Akuisisi Data (selesai unduh ERA5 dan PM2.5).\n- Tahap 2 (Bulan 2): Data Preprocessing, Data Cleansing, dan Feature Engineering.\n- Tahap 3 (Bulan 3-4): Eksperimen Pemodelan Machine Learning dan Deep Learning.\n- Tahap 4 (Bulan 5): Evaluasi Model (RMSE, MAE, R²) dan Analisis Interpretabilitas (XAI).\n- Tahap 5 (Bulan 6): Penulisan Laporan Akhir, Visualisasi Hasil, dan Presentasi Riset.')

    # 8. Hasil yang Dapat Diukur
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nHasil yang Dapat Diukur\n').bold = True
    p.add_run('- Ketersediaan dataset bersih (clean dataset) ERA5 dan PM2.5 untuk wilayah Bandung (2022-2026).\n- Performa model dengan akurasi yang lebih tinggi dan tingkat eror yang rendah (Target R² > 0.75, RMSE diminimalkan) jika dibandingkan dengan base model regresi biasa.\n- Dihasilkannya plot interpretasi (SHAP summary plot) yang memetakan fitur meteorologis terpenting yang memengaruhi tingkat PM2.5.')

    # 9. Kebaruan / Novelty
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('\nKebaruan / Novelty\n').bold = True
    p.add_run('- Pemanfaatan Variabel Lengkap ERA5 di Daerah Tropis: Mengikutsertakan variabel dinamis atmosfer atas (lapisan batas dan ozon kolom) yang jarang diekstraksi dalam riset iklim skala kota di Indonesia.\n- Model Machine Learning yang Interpretable (XAI): Kombinasi deep learning untuk mendapatkan akurasi maksimal dengan penambahan XAI, sehingga tidak hanya "bisa memprediksi" tetapi "bisa menjelaskan" mengapa suatu prediksi polusi terjadi.')

    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

if __name__ == '__main__':
    out_file = os.path.join(r'D:\Kuliah Praktik\KP BRIN\reports', 'Laporan_Presentasi_ERA5_PM25.docx')
    generate_report(out_file)
    print('Report saved to', out_file)
