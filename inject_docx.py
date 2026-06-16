import docx
from docx.shared import RGBColor, Pt
import json
import re

# Data sintesis
data_json = """[
  {"no": 1, "judul": "Improving PM2.5 Forecast Using SETAR-Tree", "metode": "SETAR-Tree & LSTM", "dataset": "PM2.5 Harian (Jan 22 - Nov 23) di Lubang Buaya, Jakarta", "hasil": "SETAR-Tree (RMSE 0.2159, MAPE 2.98% out-sample)", "kelebihan": "Sangat baik menangkap pola non-linear dan lonjakan mendadak dengan tuning minim.", "kekurangan": "Tidak memasukkan variabel cuaca/lalu lintas; kurang tangguh pada ekstrem fluktuasi puncak mendadak."},
  {"no": 2, "judul": "Comparison of LSTM and BiLSTM", "metode": "LSTM & BiLSTM", "dataset": "PM10, SO2, CO, Cuaca di Jakarta (2010-2021)", "hasil": "Baik untuk PM10 & Cuaca. Tidak ada beda signifikan LSTM vs BiLSTM.", "kelebihan": "Mengevaluasi kondisi sebelum & saat pandemi COVID-19.", "kekurangan": "Gagal memprediksi SO2, hujan, dan durasi matahari secara akurat. Anomali pandemi menurunkan akurasi."},
  {"no": 3, "judul": "LSTM Transboundary Haze SE Asia", "metode": "XGBoost (imputasi) + Bi-LSTM", "dataset": "Data kualitas udara per jam Brunei, Singapura, Thailand", "hasil": "Bi-LSTM terbaik. RMSE 9.76 - 10.96 di data uji.", "kelebihan": "XGBoost efektif mengisi data bolong stasiun cuaca, model mampu prediksi lintas negara.", "kekurangan": "Terjadi overfitting saat testing. Teknik imputasi belum sempurna untuk data lintas batas."},
  {"no": 4, "judul": "Assessing low-cost sensor PM2.5", "metode": "Regresi Linear, Uji Statistik", "dataset": "PM2.5, Kelembaban (RH), Suhu di Tamansari, Bandung (1 Tahun)", "hasil": "Korelasi linear sangat kuat dengan SuperSASS (R2 = 0.96)", "kelebihan": "Low-cost sensor terbukti tangguh untuk pemantauan real-time jarak jauh secara murah.", "kekurangan": "Alat cenderung overestimate (rata-rata 24% lebih tinggi) karena terpengaruh kelembaban udara (RH)."},
  {"no": 5, "judul": "SHAP Interpretation of XGBoost", "metode": "XGBoost + Grid Search + SHAP", "dataset": "ISPU harian (DKI Jakarta) Jan 2021 - Mei 2024", "hasil": "MAPE sangat rendah (4.44%). Parameter 'max' dan 'PM10' paling dominan.", "kelebihan": "SHAP berhasil memecahkan kelemahan black-box, memberi transparansi pengaruh variabel.", "kekurangan": "Hanya memakai basis data ISPU dasar, disarankan menambah variabel iklim dan mobilitas."},
  {"no": 6, "judul": "Deep Learning for AQ Forecasts Review", "metode": "Literature Review (CNN, RNN, STDL)", "dataset": "Skala global (Satelit, Sensor Tanah, CTMs)", "hasil": "Arsitektur Spatiotemporal (CNN-LSTM) lebih superior dibanding ML klasik.", "kelebihan": "Menyoroti keunggulan AI dibanding chemistry-transport models fisik yang lambat.", "kekurangan": "Sifat black-box AI sulit dijelaskan. Prediksi kadang tidak mematuhi kaidah fisika/kimia."},
  {"no": 7, "judul": "LSTM neural network Method dev. & eval", "metode": "LSTME (Extended LSTM)", "dataset": "PM2.5 & Cuaca per jam di 12 stasiun Beijing (2014-2016)", "hasil": "MAPE 11.93% (1 jam ke depan), Akurasi kategori AQI 90.3%", "kelebihan": "Integrasi data auxiliary (waktu & cuaca) di fully-connected layer mendongkrak performa.", "kekurangan": "Akurasi menurun drastis untuk prediksi jangka menengah/panjang (13-24 jam)."},
  {"no": 8, "judul": "Bootstrap-XGBoost & Ordinal Models", "metode": "Bootstrap, XGBoost, Ordinal Logit", "dataset": "AQI harian di Xi'an, China (Okt 22 - Sep 23)", "hasil": "XGBoost (MAPE 3.90%, R2 0.943). Ordinal Logit akurasi ranking 89.04%", "kelebihan": "Bisa mengestimasi ketidakpastian interval lewat Bootstrap.", "kekurangan": "Data latih terlalu singkat (1 tahun) sehingga tidak menangkap variasi iklim tahunan."},
  {"no": 9, "judul": "CNN-LSTM PM2.5 Beijing-Tianjin-Hebei", "metode": "Hybrid CNN-LSTM", "dataset": "PM2.5 & Faktor meteorologi per jam di 10 kota China (2021-2022)", "hasil": "CNN-LSTM menurunkan RMSE 14.30% dibanding LSTM tunggal.", "kelebihan": "Integrasi spasial dan temporal sangat stabil secara adaptif di wilayah besar.", "kekurangan": "Kualitas menurun jauh bila forward forecasting panjang. Kesulitan prediksi anomali."},
  {"no": 10, "judul": "Hybrid Spatiotemporal CNN-LSTM", "metode": "CNN, LSTM, Genetic Algorithm (GA)", "dataset": "PM2.5 di 12 Stasiun Beijing (2013-2017)", "hasil": "Ansambel model terbaik mencapai R2 0.935, RMSE 20.739.", "kelebihan": "GA mengotomatisasi pencarian parameter & seleksi fitur; memetakan stasiun 2D.", "kekurangan": "Sangat memakan waktu komputasi (computationally intensive) akibat algoritma genetik."},
  {"no": 11, "judul": "Statistical & ML Models Review", "metode": "Systematic Literature Review", "dataset": "Tinjauan global atas 412 artikel (2016-2025)", "hasil": "Random Forest (45.6%) dan LSTM (17.2%) dominan.", "kelebihan": "Memetakan tren prediktor utama global; Menyoroti absennya studi di Amerika Latin.", "kekurangan": "Mayoritas penelitian ML berstatus black-box, belum ramah policy-maker."},
  {"no": 12, "judul": "Interpretable ML Approaches Review", "metode": "Systematic Literature Review", "dataset": "56 studi global (50% China, 10% AS) periode 2011-2023", "hasil": "SHAP mendominasi metode interpretasi (46.4%), disusul PDP (17.4%).", "kelebihan": "Review komprehensif pertama yang mengupas model interpretasi untuk kualitas udara.", "kekurangan": "Mengesampingkan gray literature; banyak penelitian cacat menangani missing values."},
  {"no": 13, "judul": "Interpretable ML framework for AQ", "metode": "7 ML Algorithms (ETR, XGBoost) + SHAP", "dataset": "PM10, PM2.5, O3, Cuaca di Istanbul, Turki (2021-2023)", "hasil": "ETR terbaik di musim dingin (R2=0.93). MLP unggul PM2.5 musim panas.", "kelebihan": "Memodelkan fluktuasi multi-polutan lintas lanskap kota beda topografi.", "kekurangan": "Generalisasi spasial terbatas 3 stasiun. Hasil SHAP abai pada musim transisi."},
  {"no": 14, "judul": "Unified Approach to Interpreting Model", "metode": "SHAP (SHapley Additive exPlanations)", "dataset": "Dataset klasifikasi & regresi benchmark", "hasil": "Estimasi SHAP selaras dengan intuisi manusia dan konsisten.", "kelebihan": "Mengusulkan framework teori permainan yang sangat solid untuk XAI.", "kekurangan": "Kompleksitas komputasi NP-Hard; butuh teknik aproksimasi asumsi independen."},
  {"no": 15, "judul": "Spatio-Temporal Dynamics Extreme PM2.5", "metode": "MLR, Random Forest, XGBoost", "dataset": "PM2.5 & Cuaca (ERA5) per jam di Jakarta, Semarang, Malang", "hasil": "XGBoost mengungguli Regresi Linear (R2 = 0.534 di Jakarta).", "kelebihan": "Fokus pada iklim tropis dan manfaatkan data ERA5 terbuka untuk daerah minim data.", "kekurangan": "R2 moderat membuktikan bahwa faktor cuaca saja tidak cukup tanpa data emisi."},
  {"no": 16, "judul": "Diurnal Variations & Wavelet Coherence", "metode": "Multiple Wavelet Coherence & HYSPLIT", "dataset": "PM2.5 per setengah jam & Iklim di 7 kota Indonesia (2021)", "hasil": "Kombinasi Suhu+Kelembapan+Angin+Hujan menyetir PM2.5.", "kelebihan": "Mendeteksi korelasi berbasis frekuensi wavelet; Melacak sumber polusi gambut.", "kekurangan": "Data hanya 1 tahun anomali pandemi (2021), kurang sumber emisi lokal."},
  {"no": 17, "judul": "Estimation of PM2.5 in Jakarta", "metode": "Random Forest + GEE (Sentinel-5P)", "dataset": "Jakarta, observasi permukaan & satelit (2021-2023)", "hasil": "R2 sebesar 0.793 dengan RMSE 7.83 ug/m3.", "kelebihan": "Memetakan PM2.5 spasial secara luas mengatasi minimnya stasiun darat.", "kekurangan": "Baris dengan missing values dihapus paksa, merugikan informasi waktu."},
  {"no": 18, "judul": "Correlation of Meteorological Factors", "metode": "MLR & Korelasi Spearman", "dataset": "5 stasiun Jakarta (2022-2023)", "hasil": "MLR sangat presisi menangkap pola (R2 = 0.99).", "kelebihan": "Akurat melacak relasi linear dampak kejadian seperti El-Nino.", "kekurangan": "Nilai error aktual (RMSE) sangat masif, butuh model non-linear (ML)."},
  {"no": 19, "judul": "Sistem Prediksi LSTM Yogyakarta", "metode": "LSTM (Deep Learning)", "dataset": "7 polutan historis DLH Yogyakarta (2022-2024)", "hasil": "MAE (4.85) lebih kecil dibanding deviasi historis (10.19).", "kelebihan": "Luaran berupa purwarupa GUI klasifikasi peringatan bagi pengguna.", "kekurangan": "Sulit memprediksi HC; GUI belum bersifat live input."},
  {"no": 20, "judul": "ML for Air Quality Review", "metode": "Systematic Literature Review", "dataset": "78 publikasi global (2017-2024)", "hasil": "Model hybrid spatio-temporal catat kinerja tertinggi.", "kelebihan": "Komparasi holistik use-cases algoritma beserta pentingnya implementasi XAI.", "kekurangan": "Tantangan berupa komputasi tinggi, minim sensor lapangan, dan overfitting."},
  {"no": 21, "judul": "Framework Air Pollution Smart Cities", "metode": "WSN & IoT Architecture", "dataset": "Konsep/Desain Arsitektur Sistem Jaringan Sensor", "hasil": "Desain hemat daya dan latensi rendah.", "kelebihan": "Sistem alerting real-time bagi pemerintah kota yang hemat biaya.", "kekurangan": "Kurangnya empiris lapangan; tidak ada studi model datanya."},
  {"no": 22, "judul": "Air Quality Prediction from Images", "metode": "CNN (AQI-Net) + Grad-CAM", "dataset": "11.000 citra smartphone (Jakarta, Malang, Semarang)", "hasil": "Akurasi 99.81% dengan pelatihan sangat efisien.", "kelebihan": "Pendekatan XAI visual (Grad-CAM) untuk interpretasi prediksi CNN.", "kekurangan": "Hanya valid siang hari berpencahayaan baik; Kadang menyoroti piksel gedung."}
]"""

data = json.loads(data_json)

def set_text_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.highlight_color = None
    # Menghilangkan underline gaya hyperlink jika ada
    run.font.underline = False

doc_path = r"D:\Kuliah Praktik\KP BRIN\docs\Review_GapAnalysis_AirQuality.docx"
try:
    doc = docx.Document(doc_path)
except Exception as e:
    print("Gagal membuka dokumen:", e)
    exit(1)

# Mencari tabel sintesis (Tabel yang memiliki kolom 'Metode' atau 'Kekurangan' atau setidaknya tabel yang cukup besar)
# Di dalam template DOCX, biasanya tabel sintesis adalah tabel ke-2 atau ke-3.
target_table = None
for table in doc.tables:
    # Periksa header baris pertama
    if len(table.rows) > 0:
        row_text = "".join(cell.text.lower() for cell in table.rows[0].cells)
        if 'metode' in row_text or 'hasil' in row_text or 'kekurangan' in row_text or 'limitasi' in row_text or 'judul' in row_text:
            target_table = table
            break

if not target_table:
    print("Tabel sintesis tidak ditemukan. Menambahkan tabel baru di akhir dokumen.")
    target_table = doc.add_table(rows=1, cols=7)
    target_table.style = 'Table Grid'
    hdr_cells = target_table.rows[0].cells
    headers = ["No", "Judul", "Metode", "Dataset / Lokasi", "Hasil Terbaik", "Kelebihan", "Kekurangan / Limitasi"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                set_text_black(run)
else:
    # Kosongkan baris setelah header jika ada baris fiktif
    while len(target_table.rows) > 1:
        # docx table doesn't have a direct delete row method that's safe, 
        # but we can use XML manipulation to remove it
        tr = target_table.rows[1]._tr
        target_table._tbl.remove(tr)

# Memasukkan data ke tabel
for item in data:
    row_cells = target_table.add_row().cells
    
    texts = [
        str(item["no"]),
        item["judul"],
        item["metode"],
        item["dataset"],
        item["hasil"],
        item["kelebihan"],
        item["kekurangan"]
    ]
    
    for i, text in enumerate(texts):
        # We assign text directly, then clear existing runs and add a new run to ensure strict formatting
        # Some tables might have different column counts, so we check boundaries
        if i < len(row_cells):
            cell = row_cells[i]
            cell.text = "" # Bersihkan teks default
            p = cell.paragraphs[0]
            run = p.add_run(text)
            # Set font color to absolute black
            set_text_black(run)
            
            # Hapus warna background sel jika ada (agar tidak putih mentereng jika template punya)
            # Ini memerlukan modifikasi properti XML shd, tapi kita biarkan transparan (default).

# Selain itu, ubah seluruh teks di dokumen agar font color-nya menjadi hitam jika sebelumnya ada yang biru (hyperlink)
for p in doc.paragraphs:
    for run in p.runs:
        # Jika font biru atau ada tema biru, ganti hitam
        if run.font.color and run.font.color.rgb:
            if run.font.color.rgb != RGBColor(0,0,0):
                set_text_black(run)
        else:
            # Paksa hitam
            set_text_black(run)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.color and run.font.color.rgb:
                        if run.font.color.rgb != RGBColor(0,0,0):
                            set_text_black(run)
                    else:
                        set_text_black(run)

doc.save(doc_path.replace('.docx', '_Filled.docx'))
print("Berhasil memasukkan data sintesis ke tabel DOCX dan mengubah warna teks menjadi hitam murni. File disimpan sebagai Review_GapAnalysis_AirQuality_Filled.docx.")
