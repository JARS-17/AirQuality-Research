import docx
from docx.shared import Pt, RGBColor

doc_path = r"D:\Kuliah Praktik\KP BRIN\docs\Review_GapAnalysis_AirQuality_Filled.docx"
try:
    doc = docx.Document(doc_path)
except Exception as e:
    print("Gagal membuka dokumen:", e)
    exit(1)

def apply_formatting(run):
    run.font.name = "Times New Roman"
    if run.font.size is None:
        run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.highlight_color = None

# Terapkan ke semua paragraf utama
for p in doc.paragraphs:
    for run in p.runs:
        apply_formatting(run)

# Terapkan ke semua tabel
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    apply_formatting(run)

# Simpan kembali
doc.save(doc_path)
print("Seluruh dokumen berhasil dinormalisasi: Times New Roman, ukuran 12, warna Hitam murni.")
